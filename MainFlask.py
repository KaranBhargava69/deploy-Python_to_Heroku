from flask import Flask, request, jsonify
import os
import requests
from docx import Document
from docx2pdf import convert
from datetime import datetime

app = Flask(__name__)

# Your Salesforce authentication credentials and other global variables
client_id = '3MVG9wt4IL4O5wvIiODM3C3u77INsVu0YYj1Ews1FZM6IgZ0Ees8AATpMsovD.S6MRA_0zDv5iAUU_lZpARB.'
client_secret = 'FD73D5AC526D50D88E8FFA2F3ED88797741F2F496782CDFAAE25437421CD5F91'
username = 'scheduledsubflow@karan.com'
password = 'subflow@karan.com911'
security_token = 'DZebWiyVvh2zHLtjQfA7IW5Y'
login_url = 'https://login.salesforce.com/services/oauth2/token'

@app.route('/fetch_data', methods=['GET'])
def fetch_data():
    record_id = request.args.get('record_id')  # Get the record ID from the query parameter
    # Salesforce authentication payload
    payload = {
        'grant_type': 'password',
        'client_id': client_id,
        'client_secret': client_secret,
        'username': username,
        'password': password + security_token
    }
    
    today_date = datetime.now()

    formatted_date = today_date.strftime("%B %d, %Y")

    # Authenticate and get access token
    response = requests.post(login_url, data=payload)
    if response.status_code == 200:
        access_token = response.json()['access_token']
        instance_url = response.json()['instance_url']

        # Modify the query to fetch data based on the record ID
        query_endpoint = instance_url + f'/services/data/v50.0/query/?q=SELECT+Id,Name+FROM+Opportunity+WHERE+Id=\'{record_id}\''
#         query_endpoint = instance_url + '/services/data/v50.0/query/?q=SELECT+Id,Name+FROM+Opportunity'

        # Make request to fetch data
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': "application/json"
        }

        query_response = requests.get(query_endpoint, headers=headers)

        if query_response.status_code == 200:
            salesforce_data = query_response.json()
            # Open the Word document
            file_path = 'D:\\New folder\\ConAm Internal\\Managed Service Temp.docx'
            doc = Document(file_path)

            placeholders = {
                    '{Name}': salesforce_data['records'][0]['Name'],
                    '{today’sdate}':formatted_date,
                    #'{Id}': salesforce_data['records'][0]['Id'],
                    # Add more placeholders as needed
                }
            # Define a function to replace placeholders in different parts of the document
            def replace_placeholders(doc_part):
                for paragraph in doc_part.paragraphs:
                    for placeholder, value in placeholders.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

                for table in doc_part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for placeholder, value in placeholders.items():
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, value)

            # Replace placeholders in the main body
            replace_placeholders(doc)

            # Replace placeholders in headers
            for section in doc.sections:
                for header in section.header.paragraphs:
                    for paragraph in header.runs:
                        for placeholder, value in placeholders.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)

            # Replace placeholders in footers
            for section in doc.sections:
                for footer in section.footer.paragraphs:
                    for paragraph in footer.runs:
                        for placeholder, value in placeholders.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)


            for section in doc.sections:
                header = section.first_page_header# Access the header of the current section
                if header is not None: # Check if a header exists
                    for paragraph in header.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, value in placeholders.items():     
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)


            # Save the modified document temporarily
            temp_file_path = os.path.splitext(file_path)[0] + '_modified.docx'
            doc.save(temp_file_path)

            # Convert the modified docx to PDF directly
            convert(temp_file_path)

            # Remove the temporary modified docx file
            os.remove(temp_file_path)

            return jsonify({"message": "Document generated successfully."})
        else:
            return jsonify({"error": "Failed to fetch data from Salesforce."}), 500
    else:
        return jsonify({"error": "Failed to authenticate with Salesforce."}), 500

if __name__ == '__main__':
    app.run(port=8080)

